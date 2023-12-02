import tkinter
from tkinter import *
import tkinter as tk
import tkinter.messagebox as mbox
from PIL import Image, ImageTk
from fpdf import FPDF
import webbrowser
from transformers import pipeline
import spacy
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT 

selected_option = None
# main window created
window = Tk()
window.geometry("1000x700")
window.title("Resume Generator")

def def_start():
    nlp_ner = spacy.load("model-best\content\model-last")

    def def_PDF():
        global basic
        basic = text_enter.get("1.0", END)
        text_generator = pipeline("text-generation", model="gpt2")

        mbox.showinfo("Generating Resume", "Resume generation in progress...")

        generated_text = text_generator(basic, max_length=100, num_return_sequences=1)

        doc = nlp_ner(generated_text[0]['generated_text'])
        recognized_entities = [(ent.text, ent.label_) for ent in doc.ents]

        # Create a new Word document
        resume_doc = Document()

        # Initialize a variable to keep track of the current label
        current_label = None
        name_label_detected = False

        for text, label in recognized_entities:
            print(label)
            if label == "Name":
                print("what")
            # For the first occurrence of 'Name', make the text bigger and centered without subheading
                paragraph = resume_doc.add_paragraph()
                print("hello")
                run = paragraph.add_run(text)
                font = run.font
                font.size = Pt(20)  # Font size for 'Name' (adjust as needed)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align 'Name'
                name_label_detected = True  # Se

            elif label != current_label:
                print()
                # If the label changes, create a new heading and set it as the current label
                current_label = label
                resume_doc.add_heading(label, level=1)
                resume_doc.add_paragraph(text)
            else:
                # If the label is the same as the current one, add content under the same heading
                resume_doc.add_paragraph(text)

        # Add a custom style for the 'Name' label (bigger font size and centered)

        # Save the document
        resume_doc.save("generated_resume.docx")

        mbox.showinfo("Resume Generated", "Resume has been generated successfully!")


    f1 = Frame(window, width=1000, height=700)
    f1.propagate(0)
    f1.pack(side='top')

    start1 = tk.Label(f1,text="RESUME GENERATOR", font=("Arial", 50), fg="magenta")  # same way bg
    start1.place(x=120, y=10)

    text_enter = tk.Text(window, height=13, width=48, font=("Arial", 20), bg="light blue", fg="blue", borderwidth=3,relief="solid")
    text_enter.place(x=120, y=110)

    def clear_text():
        text_enter.delete("1.0", END)

        # created button for clear
    clearb = Button(window, text="CLEAR", command=clear_text, font=("Arial", 25), bg="light green", fg="blue",borderwidth=3, relief="raised")
    clearb.place(x=100, y=580)

    pdfb = Button(window, text="GENERATE RESUME", command=def_PDF, font=("Arial", 25), bg="orange", fg="blue",borderwidth=3, relief="raised")
    pdfb.place(x=300, y=590)

start1 = tk.Label(text = "RESUME GENERATOR", font=("Arial", 55), fg="magenta") # same way bg
start1.place(x = 90, y = 10)

# image on the main window
path = "Images/resume_front.png"
# Creates a Tkinter-compatible photo image, which can be used everywhere Tkinter expects an image object.
img = ImageTk.PhotoImage(Image.open(path))
# The Label widget is a standard Tkinter widget used to display a text or image on the screen.
panel = tk.Label(window, image = img)
panel.place(x = 340, y = 150)

# created start button
startb = Button(window, text="START",command=def_start,font=("Arial", 30), bg = "light green", fg = "blue", borderwidth=3, relief="raised")
startb.place(x =100 , y =570 )

#Define a callback function
def callback(url):
   webbrowser.open_new_tab(url)

# created EMI button
templateb = Button(window, text="Analyser",command=lambda:callback("https://novoresume.com/resume-templates"),font=("Arial", 30), bg = "orange", fg = "blue", borderwidth=3, relief="raised")
templateb.place(x =370 , y =570 )

# function for exiting
def exit_win():
    if mbox.askokcancel("Exit", "Do you want to exit?"):
        window.destroy()

# created exit button
exitb = Button(window, text="EXIT",command=exit_win,font=("Arial", 30), bg = "red", fg = "blue", borderwidth=3, relief="raised")
exitb.place(x =730 , y =570 )


window.protocol("WM_DELETE_WINDOW", exit_win)
window.mainloop()
