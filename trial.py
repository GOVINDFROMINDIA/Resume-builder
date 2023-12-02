
# Resume Generator

# imported necessary library
import tkinter
from tkinter import *
import tkinter as tk
import tkinter.messagebox as mbox
from PIL import Image, ImageTk
from fpdf import FPDF
import webbrowser
from transformers import pipeline
import spacy
import json
from docx import Document

selected_option = None
# main window created
window = Tk()
window.geometry("1000x700")
window.title("Resume Generator")


# defined function for start button
def def_start():
    nlp_ner = spacy.load("model-best\content\model-last")
    def basic_fun():
        def enter_fun():
            global basic
            basic = text_enter.get("1.0", END)
            window1.destroy()

        # created main window
        window1 = Tk()
        window1.geometry("1000x700")
        window1.title("BASIC DETAILS")

        # for top label
        top1 = tk.Label(window1, text="Basic Info", font=("Arial", 50), fg="green")  # same way bg
        top1.place(x=330, y=10)

        # Taking input of code from TextArea
        text_enter = tk.Text(window1, height=13, width=48, font=("Arial", 20), bg="light blue", fg="blue", borderwidth=3,relief="solid")
        text_enter.place(x=120, y=110)

        def clear_text():
            text_enter.delete("1.0", END)

        # created button for clear
        clearb = Button(window1, text="CLEAR", command=clear_text, font=("Arial", 25), bg="light green", fg="blue",borderwidth=3, relief="raised")
        clearb.place(x=200, y=580)

        # created button for enter
        enterb = Button(window1, text="ENTER", command=enter_fun, font=("Arial", 25), bg="light green", fg="blue",borderwidth=3, relief="raised")
        enterb.place(x=650, y=580)

    def edu_fun():
        def enter_fun():
            global edu
            edu = text_enter.get("1.0", END)
            window1.destroy()

        # created main window
        window1 = Tk()
        window1.geometry("1000x700")
        window1.title("EDUCATION DETAILS")

        # for top label
        top1 = tk.Label(window1, text="Education", font=("Arial", 50), fg="green")  # same way bg
        top1.place(x=340, y=10)

        # Taking input of code from TextArea
        text_enter = tk.Text(window1, height=13, width=48, font=("Arial", 20), bg="light blue", fg="blue",
                             borderwidth=3, relief="solid")
        text_enter.place(x=120, y=110)

        def clear_text():
            text_enter.delete("1.0", END)

        # created button for clear
        clearb = Button(window1, text="CLEAR", command=clear_text, font=("Arial", 25), bg="light green", fg="blue",
                        borderwidth=3, relief="raised")
        clearb.place(x=200, y=580)

        # created button for enter
        enterb = Button(window1, text="ENTER", command=enter_fun, font=("Arial", 25), bg="light green", fg="blue",
                        borderwidth=3, relief="raised")
        enterb.place(x=650, y=580)

    def skill_fun():
        def enter_fun():
            global skill
            skill = text_enter.get("1.0", END)
            window1.destroy()

        # created main window
        window1 = Tk()
        window1.geometry("1000x700")
        window1.title("SKILLS DETAILS")

        # for top label
        top1 = tk.Label(window1, text="Skills", font=("Arial", 50), fg="green")  # same way bg
        top1.place(x=390, y=10)

        # Taking input of code from TextArea
        text_enter = tk.Text(window1, height=13, width=48, font=("Arial", 20), bg="light blue", fg="blue",
                             borderwidth=3, relief="solid")
        text_enter.place(x=120, y=110)

        def clear_text():
            text_enter.delete("1.0", END)

        # created button for clear
        clearb = Button(window1, text="CLEAR", command=clear_text, font=("Arial", 25), bg="light green", fg="blue",
                        borderwidth=3, relief="raised")
        clearb.place(x=200, y=580)

        # created button for enter
        enterb = Button(window1, text="ENTER", command=enter_fun, font=("Arial", 25), bg="light green", fg="blue",
                        borderwidth=3, relief="raised")
        enterb.place(x=650, y=580)

    def exp_fun():
        def enter_fun():
            global exp
            exp = text_enter.get("1.0", END)
            window1.destroy()

        # created main window
        window1 = Tk()
        window1.geometry("1000x700")
        window1.title("EXPERIENCE DETAILS")

        # for top label
        top1 = tk.Label(window1, text="Experience", font=("Arial", 50), fg="green")  # same way bg
        top1.place(x=340, y=10)

        # Taking input of code from TextArea
        text_enter = tk.Text(window1, height=13, width=48, font=("Arial", 20), bg="light blue", fg="blue",
                             borderwidth=3, relief="solid")
        text_enter.place(x=120, y=110)

        def clear_text():
            text_enter.delete("1.0", END)

        # created button for clear
        clearb = Button(window1, text="CLEAR", command=clear_text, font=("Arial", 25), bg="light green", fg="blue",
                        borderwidth=3, relief="raised")
        clearb.place(x=200, y=580)

        # created button for enter
        enterb = Button(window1, text="ENTER", command=enter_fun, font=("Arial", 25), bg="light green", fg="blue",
                        borderwidth=3, relief="raised")
        enterb.place(x=650, y=580)

    def canda_fun():
        def enter_fun():
            global canda
            canda = text_enter.get("1.0", END)
            window1.destroy()

        # created main window
        window1 = Tk()
        window1.geometry("1000x700")
        window1.title("ADDITIONAL INFO")

        # for top label
        top1 = tk.Label(window1, text="Additional INFO", font=("Arial", 50), fg="green")  # same way bg
        top1.place(x=130, y=10)

        # Taking input of code from TextArea
        text_enter = tk.Text(window1, height=13, width=48, font=("Arial", 20), bg="light blue", fg="blue",
                             borderwidth=3, relief="solid")
        text_enter.place(x=120, y=110)

        def clear_text():
            text_enter.delete("1.0", END)

        # created button for clear
        clearb = Button(window1, text="CLEAR", command=clear_text, font=("Arial", 25), bg="light green", fg="blue",
                        borderwidth=3, relief="raised")
        clearb.place(x=200, y=580)

        # created button for enter
        enterb = Button(window1, text="ENTER", command=enter_fun, font=("Arial", 25), bg="light green", fg="blue",
                        borderwidth=3, relief="raised")
        enterb.place(x=650, y=580)
    
    def achievement_fun():
        def enter_fun():
            global awards
            awards = text_enter.get("1.0", END)
            window1.destroy()

        # created main window
        window1 = Tk()
        window1.geometry("1000x700")
        window1.title("ADDITIONAL INFO")

        # for top label
        top1 = tk.Label(window1, text="Achievements", font=("Arial", 50), fg="green")  # same way bg
        top1.place(x=130, y=10)

        # Taking input of code from TextArea
        text_enter = tk.Text(window1, height=13, width=48, font=("Arial", 20), bg="light blue", fg="blue",
                             borderwidth=3, relief="solid")
        text_enter.place(x=120, y=110)

        def clear_text():
            text_enter.delete("1.0", END)

        # created button for clear
        clearb = Button(window1, text="CLEAR", command=clear_text, font=("Arial", 25), bg="light green", fg="blue",
                        borderwidth=3, relief="raised")
        clearb.place(x=200, y=580)

        # created button for enter
        enterb = Button(window1, text="ENTER", command=enter_fun, font=("Arial", 25), bg="light green", fg="blue",
                        borderwidth=3, relief="raised")
        enterb.place(x=650, y=580)
    def add_section():
        global add_flag,selected_option
        add_flag = True
        mbox.showinfo("Note", "You can add the optional section now.")
        options = ["Additional Info", "Achievements"]
        selected_option = tk.StringVar()
        dropdown = tk.OptionMenu(window, selected_option, *options)
        dropdown.config(font=("Arial", 17), bg="light green", fg="blue", borderwidth=3, relief="raised")
        dropdown.place(x=700, y=500)
    # Binding the function to handle dropdown selection
        selected_option.trace("w", handle_selection)

    def handle_selection(*args):
        global selected_option
        selected = selected_option.get()
        if selected == "Additional Info":
            canda1 = tk.Label(f1, text="Enter Additional Info  :  ", font=("Arial", 30), fg="green")  # same way bg
            canda1.place(x=100, y=350)
            candab = Button(window, text="HERE", command=canda_fun, font=("Arial", 17), bg="light green", fg="blue", borderwidth=3, relief="raised")
            candab.place(x=800, y=400)
        elif selected == "Achievements":
            canda1 = tk.Label(f1, text="Enter Achievements  :  ", font=("Arial", 30), fg="green")  # same way bg
            canda1.place(x=100, y=380)
            candab = Button(window, text="HERE", command=achievement_fun, font=("Arial", 17), bg="light green", fg="blue", borderwidth=3,relief="raised")
            candab.place(x=800, y=380)

    # function for generating and saving the PDF
    def def_PDF():

        combined_data = f"{basic} {edu} {skill} {exp} {canda} {awards}"
        text_generator = pipeline("text-generation", model="gpt2")

        if not (basic and edu and skill and exp) and not add_flag:
            mbox.showwarning("Warning", "Please fill all mandatory sections!")
        else:
            mbox.showinfo("Generating Resume", "Resume generation in progress...")

            generated_text = text_generator(combined_data, max_length=100, num_return_sequences=1)

            doc = nlp_ner(generated_text[0]['generated_text'])
            recognized_entities = [(ent.text, ent.label_) for ent in doc.ents]

            # Create a new Word document
            resume_doc = Document()

            # Add recognized labels as headings and corresponding content to the document
            for text, label in recognized_entities:
                resume_doc.add_heading(label, level=1)
                resume_doc.add_paragraph(text)

            # Save the document
            resume_doc.save("generated_resume.docx")

            mbox.showinfo("Resume Generated", "Resume has been generated successfully!")

    # new frame created
    f1 = Frame(window, width=1000, height=700)
    f1.propagate(0)
    f1.pack(side='top')

    # for top label
    start1 = tk.Label(f1,text="RESUME GENERATOR", font=("Arial", 50), fg="magenta")  # same way bg
    start1.place(x=120, y=10)
    # for education label
    edu1 = tk.Label(f1, text="Enter Your Education  :  ", font=("Arial", 30), fg="green")  # same way bg
    edu1.place(x=100, y=160)
    # created button for education
    edub = Button(window, text="HERE", command=edu_fun, font=("Arial", 17), bg="light green", fg="blue", borderwidth=3,relief="raised")
    edub.place(x=800, y=160)

    # for skill label
    skill1 = tk.Label(f1, text="Enter Your Skills  :  ", font=("Arial", 30), fg="green")  # same way bg
    skill1.place(x=100, y=220)
    # created button for skill
    skillb = Button(window, text="HERE", command=skill_fun, font=("Arial", 17), bg="light green", fg="blue", borderwidth=3,relief="raised")
    skillb.place(x=800, y=220)

    # for Experience label
    exp1 = tk.Label(f1, text="Enter Your Experience  :  ", font=("Arial", 30), fg="green")  # same way bg
    exp1.place(x=100, y=280)
    # created button for Experience
    expb = Button(window, text="HERE", command=exp_fun, font=("Arial", 17), bg="light green", fg="blue", borderwidth=3,relief="raised")
    expb.place(x=800, y=280)

    candab = Button(window, text="HERE", command=canda_fun, font=("Arial", 17), bg="light green", fg="blue", borderwidth=3,relief="raised")
    candab.place(x=800, y=350)
    basic1 = tk.Label(f1, text="Enter Your Designation/About  :  ", font=("Arial", 30), fg="green")  # same way bg
    basic1.place(x=100, y=100)
    basicb = Button(window, text="HERE", command=basic_fun, font=("Arial", 17), bg="light green", fg="blue",borderwidth=3, relief="raised")
    basicb.place(x=800, y=100)

    add_section_btn = Button(window, text="Add Section", command=add_section, font=("Arial", 17), bg="light green", fg="blue", borderwidth=3, relief="raised")
    add_section_btn.place(x=500, y=500)

    # created button for generating RESUME
    pdfb = Button(window, text="GENERATE RESUME", command=def_PDF, font=("Arial", 25), bg="orange", fg="blue",borderwidth=3, relief="raised")
    pdfb.place(x=300, y=590)


# top label
start1 = tk.Label(text = "RESUME GENERATOR", font=("Arial", 55), fg="magenta") # same way bg
start1.place(x = 90, y = 10)

# image on the main window
path = "Images/resume_front.png"
# Creates a Tkinter-compatible photo image, which can be used everywhere Tkinter expects an image object.
img = ImageTk.PhotoImage(Image.open(path))
# The Label widget is a standard Tkinter widget used to display a text or image on the screen.
panel = tk.Label(window, image = img)
panel.place(x = 340, y = 150)

# created start button
startb = Button(window, text="START",command=def_start,font=("Arial", 30), bg = "light green", fg = "blue", borderwidth=3, relief="raised")
startb.place(x =100 , y =570 )

#Define a callback function
def callback(url):
   webbrowser.open_new_tab(url)

# created EMI button
templateb = Button(window, text="Analyser",command=lambda:callback("https://novoresume.com/resume-templates"),font=("Arial", 30), bg = "orange", fg = "blue", borderwidth=3, relief="raised")
templateb.place(x =370 , y =570 )

# function for exiting
def exit_win():
    if mbox.askokcancel("Exit", "Do you want to exit?"):
        window.destroy()

# created exit button
exitb = Button(window, text="EXIT",command=exit_win,font=("Arial", 30), bg = "red", fg = "blue", borderwidth=3, relief="raised")
exitb.place(x =730 , y =570 )


window.protocol("WM_DELETE_WINDOW", exit_win)
window.mainloop()

